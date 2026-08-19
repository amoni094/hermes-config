#!/usr/bin/env python3
"""
redact_pii.py — Australian Trust Deed PII Redactor
Redacts personal information within the meaning of Privacy Act 1988 (Cth) s 6.

Usage:
    python3 redact_pii.py --input deed.pdf [--output deed_redacted.pdf]
                          [--log audit.json] [--quality-report ocr.json]
                          [--adversarial-passes 3] [--dry-run] [--strict]
                          [--check-ocr-only] [--no-ner]
                          [--min-confidence 0.6]
                          [--whitelist-entity "Smith Family Trust Pty Ltd"]
"""

from __future__ import annotations

import argparse
import json
import re
import sys
import traceback
from datetime import datetime, timezone
from pathlib import Path
from typing import Any


# ---------------------------------------------------------------------------
# Dependency check (do not auto-install)
# ---------------------------------------------------------------------------

REQUIRED_PKGS: dict[str, str] = {
    "presidio_analyzer": "presidio-analyzer",
    "presidio_anonymizer": "presidio-anonymizer",
    "fitz": "pymupdf",
    "pdfplumber": "pdfplumber",
    "docx": "python-docx",
    "PIL": "pillow",
    "spacy": "spacy",
}


def check_deps(no_ner: bool = False) -> list[str]:
    missing: list[str] = []
    for module, pkg in REQUIRED_PKGS.items():
        if no_ner and module == "spacy":
            continue
        try:
            __import__(module)
        except ImportError:
            missing.append(f"  pip install {pkg}")
    return missing


# ---------------------------------------------------------------------------
# OCR quality assessment
# ---------------------------------------------------------------------------

MIN_CHARS_PER_PAGE = 50   # pages with fewer chars are considered empty
GOOD_SCORE_THRESHOLD = 0.80
LOW_CONFIDENCE_THRESHOLD = 0.50


def assess_ocr_quality(input_path: Path) -> dict[str, Any]:
    """
    Assess the text-layer quality of a PDF.
    Returns a dict with ocr_quality, ocr_confidence_score, pages_with_text,
    total_pages, empty_pages.
    For DOCX, always returns 'good' (DOCX has no OCR layer).
    """
    if input_path.suffix.lower() in (".docx", ".doc"):
        return {
            "ocr_quality": "good",
            "ocr_confidence_score": 1.0,
            "pages_with_text": 1,
            "total_pages": 1,
            "empty_pages": [],
            "format": "docx",
        }

    import fitz  # type: ignore[import-untyped]  # pymupdf

    doc = fitz.open(str(input_path))
    total_pages: int = doc.page_count
    pages_with_text = 0
    empty_pages: list[int] = []

    for i in range(total_pages):
        page = doc[i]
        text: str = str(page.get_text("text"))
        char_count = len(text.strip())
        if char_count >= MIN_CHARS_PER_PAGE:
            pages_with_text += 1
        else:
            empty_pages.append(i + 1)  # 1-indexed for human report

    doc.close()

    score = pages_with_text / total_pages if total_pages > 0 else 0.0

    if score >= GOOD_SCORE_THRESHOLD and not empty_pages:
        quality = "good"
    elif score >= LOW_CONFIDENCE_THRESHOLD:
        quality = "low_confidence"
    else:
        quality = "unreadable"

    return {
        "ocr_quality": quality,
        "ocr_confidence_score": round(score, 4),
        "pages_with_text": pages_with_text,
        "total_pages": total_pages,
        "empty_pages": empty_pages,
        "format": "pdf",
    }


# ---------------------------------------------------------------------------
# Australian-specific Presidio custom recognisers
# ---------------------------------------------------------------------------

def build_au_recognisers() -> list[Any]:
    """Build custom AU Presidio pattern recognisers."""
    from presidio_analyzer import Pattern, PatternRecognizer  # type: ignore[import-untyped]

    recognisers: list[Any] = []

    # TFN: 8 or 9 digits optionally grouped with spaces
    recognisers.append(PatternRecognizer(
        supported_entity="AU_TFN",
        name="AU_TFN_Recogniser",
        patterns=[
            Pattern(name="tfn_spaced", regex=r"\b\d{3}[ \-]?\d{3}[ \-]?\d{2,3}\b", score=0.75),
            Pattern(name="tfn_label", regex=r"(?i)(?:TFN|Tax File Number)[:\s]+[\d\s\-]{8,11}", score=0.85),
        ],
        context=["tax file", "tfn"],
    ))

    # Medicare: 10-digit optionally spaced
    recognisers.append(PatternRecognizer(
        supported_entity="AU_MEDICARE",
        name="AU_Medicare_Recogniser",
        patterns=[
            Pattern(name="medicare_spaced", regex=r"\b\d{4}[ ]?\d{5}[ ]?\d{1}\b", score=0.80),
            Pattern(name="medicare_label", regex=r"(?i)Medicare\s+(?:Number|No\.?|Card)[:\s]+[\d ]+", score=0.90),
        ],
        context=["medicare", "health card"],
    ))

    # ABN (individual only — label-triggered; corporate ABNs whitelisted downstream)
    recognisers.append(PatternRecognizer(
        supported_entity="AU_ABN_INDIVIDUAL",
        name="AU_ABN_Individual_Recogniser",
        patterns=[
            Pattern(
                name="abn_individual_label",
                regex=r"(?i)(?:ABN|Australian Business Number)[:\s]+\d{2}[ ]?\d{3}[ ]?\d{3}[ ]?\d{3}",
                score=0.70,
            ),
        ],
        context=["abn", "australian business number"],
    ))

    # Australian driver licence
    recognisers.append(PatternRecognizer(
        supported_entity="AU_DRIVER_LICENCE",
        name="AU_DriverLicence_Recogniser",
        patterns=[
            Pattern(
                name="dl_label",
                regex=r"(?i)Driver['\s]*s?\s*Licen[cs]e\s*(?:Number|No\.?)?[:\s]+[A-Z0-9]{5,12}",
                score=0.85,
            ),
        ],
        context=["driver", "licence", "license", "dl number"],
    ))

    # Australian passport number: one letter + 7 digits
    recognisers.append(PatternRecognizer(
        supported_entity="AU_PASSPORT",
        name="AU_Passport_Recogniser",
        patterns=[
            Pattern(name="passport_pattern", regex=r"\b[A-Z][0-9]{7}\b", score=0.65),
            Pattern(
                name="passport_label",
                regex=r"(?i)Passport\s*(?:Number|No\.?)[:\s]+[A-Z][0-9]{7}",
                score=0.90,
            ),
        ],
        context=["passport", "travel document"],
    ))

    # Australian BSB + account
    recognisers.append(PatternRecognizer(
        supported_entity="AU_BANK_ACCOUNT",
        name="AU_BankAccount_Recogniser",
        patterns=[
            Pattern(name="bsb_account", regex=r"\b\d{3}[-]?\d{3}\s+\d{5,10}\b", score=0.70),
            Pattern(
                name="bsb_label",
                regex=r"(?i)(?:BSB|Bank[:\s]+)[\d\-]+\s+(?:Account|Acct)[:\s]+[\d ]+",
                score=0.85,
            ),
        ],
        context=["bsb", "account number", "bank account"],
    ))

    # Date of birth (explicit label)
    recognisers.append(PatternRecognizer(
        supported_entity="DATE_OF_BIRTH",
        name="AU_DOB_Recogniser",
        patterns=[
            Pattern(
                name="dob_label",
                regex=r"(?i)(?:Date of Birth|D\.?O\.?B\.?|Born)[:\s]+\d{1,2}[/\-\.]\d{1,2}[/\-\.]\d{2,4}",
                score=0.90,
            ),
            Pattern(
                name="dob_label_long",
                regex=r"(?i)(?:Date of Birth|D\.?O\.?B\.?|Born)[:\s]+\d{1,2}\s+\w+\s+\d{4}",
                score=0.88,
            ),
        ],
        context=["born", "dob", "date of birth"],
    ))

    # Australian residential address
    recognisers.append(PatternRecognizer(
        supported_entity="AU_ADDRESS",
        name="AU_Address_Recogniser",
        patterns=[
            Pattern(
                name="au_street",
                regex=(
                    r"\b\d{1,4}\s+[A-Z][a-zA-Z\s]{2,30}"
                    r"(?:Street|St|Avenue|Ave|Road|Rd|Drive|Dr|Court|Ct|"
                    r"Place|Pl|Lane|Ln|Crescent|Cres|Close|Cl|Way|Boulevard|Blvd|Parade|Pde|"
                    r"Circuit|Cct|Grove|Gr|Terrace|Tce|Highway|Hwy)\b"
                ),
                score=0.72,
            ),
            Pattern(
                name="au_postcode_state",
                regex=r"\b(?:NSW|VIC|QLD|SA|WA|TAS|ACT|NT)\s+\d{4}\b",
                score=0.65,
            ),
        ],
        context=["address", "residing", "resident", "lives at", "home", "postal"],
    ))

    return recognisers


# ---------------------------------------------------------------------------
# Corporate entity whitelist heuristic
# ---------------------------------------------------------------------------

CORPORATE_SUFFIXES = re.compile(
    r"\b(?:Pty\.?\s*Ltd\.?|Ltd\.?|Inc\.?|Corp\.?|LLC|"
    r"Limited|Partnership|Incorporated|Trustees?|Superannuation Fund)\b",
    re.IGNORECASE,
)


def is_likely_corporate(text: str) -> bool:
    return bool(CORPORATE_SUFFIXES.search(text))


# ---------------------------------------------------------------------------
# Trust name detection
# ---------------------------------------------------------------------------

TRUST_NAME_PATTERN = re.compile(
    r"\b[A-Z][a-zA-Z\s''\-]{2,40}(?:Family\s+)?Trust\b",
    re.IGNORECASE,
)


def extract_trust_names(full_text: str) -> list[str]:
    return list({m.group(0).strip() for m in TRUST_NAME_PATTERN.finditer(full_text)})


# ---------------------------------------------------------------------------
# PII detection via Presidio
# ---------------------------------------------------------------------------

REDACT_ENTITIES = [
    "PERSON",
    "EMAIL_ADDRESS",
    "PHONE_NUMBER",
    "AU_TFN",
    "AU_MEDICARE",
    "AU_ABN_INDIVIDUAL",
    "AU_DRIVER_LICENCE",
    "AU_PASSPORT",
    "AU_BANK_ACCOUNT",
    "DATE_OF_BIRTH",
    "AU_ADDRESS",
    "DATE_TIME",        # filtered below to DOB context only
    "MEDICAL_LICENSE",  # sensitive info
    "NRP",              # nationality / race / political opinion
]


def build_analyzer(no_ner: bool, min_confidence: float) -> Any:  # noqa: ARG001 (min_confidence reserved for future use)
    from presidio_analyzer import AnalyzerEngine  # type: ignore[import-untyped]

    if no_ner:
        engine: Any = AnalyzerEngine()
    else:
        import spacy  # type: ignore[import-untyped]
        from presidio_analyzer.nlp_engine import NlpEngineProvider  # type: ignore[import-untyped]

        model: str | None = None
        for candidate in ("en_core_web_lg", "en_core_web_sm"):
            try:
                spacy.load(candidate)
                model = candidate
                break
            except OSError:
                continue

        if model is None:
            print(
                "WARNING: No spaCy model found. Falling back to regex-only detection. "
                "Install with: python -m spacy download en_core_web_lg",
                file=sys.stderr,
            )
            engine = AnalyzerEngine()
        else:
            if model == "en_core_web_sm":
                print(
                    "WARNING: en_core_web_sm loaded (lower recall). "
                    "Install en_core_web_lg for production: python -m spacy download en_core_web_lg",
                    file=sys.stderr,
                )
            config = {
                "nlp_engine_name": "spacy",
                "models": [{"lang_code": "en", "model_name": model}],
            }
            provider = NlpEngineProvider(nlp_configuration=config)
            nlp_engine = provider.create_engine()
            engine = AnalyzerEngine(nlp_engine=nlp_engine)

    # Register AU custom recognisers
    for rec in build_au_recognisers():
        engine.registry.add_recognizer(rec)

    return engine


def analyze_text(
    analyzer: Any,
    text: str,
    min_confidence: float,
    whitelist_entities: list[str],
) -> list[dict[str, Any]]:
    results = analyzer.analyze(
        text=text,
        entities=REDACT_ENTITIES,
        language="en",
        score_threshold=min_confidence,
    )

    findings: list[dict[str, Any]] = []
    for r in results:
        matched: str = text[r.start:r.end]

        # Skip corporate names for PERSON entity
        if r.entity_type == "PERSON" and is_likely_corporate(matched):
            continue

        # Skip user-supplied whitelist
        if any(w.lower() in matched.lower() for w in whitelist_entities):
            continue

        # DATE_TIME: only include if DOB context precedes the match
        if r.entity_type == "DATE_TIME":
            context_window = text[max(0, r.start - 40): r.start].lower()
            if not any(kw in context_window for kw in ("dob", "born", "date of birth", "d.o.b")):
                continue

        findings.append({
            "category": r.entity_type,
            "matched_text": matched,
            "presidio_score": round(r.score, 4),
            "char_start": r.start,
            "char_end": r.end,
        })

    return findings


# ---------------------------------------------------------------------------
# PDF redaction
# ---------------------------------------------------------------------------

REDACT_LABEL = "[REDACTED]"


def redact_pdf(
    input_path: Path,
    output_path: Path,
    analyzer: Any,
    min_confidence: float,
    whitelist_entities: list[str],
    adversarial_passes: int,
) -> list[dict[str, Any]]:
    import fitz  # type: ignore[import-untyped]

    passes_log: list[dict[str, Any]] = []
    current_path = input_path

    for pass_num in range(1, adversarial_passes + 1):
        doc = fitz.open(str(current_path))
        all_findings: list[dict[str, Any]] = []

        for page_idx in range(doc.page_count):
            page = doc[page_idx]
            page_text: str = str(page.get_text("text"))
            findings = analyze_text(analyzer, page_text, min_confidence, whitelist_entities)

            for finding in findings:
                finding["page"] = page_idx + 1
                matched: str = finding["matched_text"]
                rects = page.search_for(matched, quads=False)
                for rect in rects:
                    page.add_redact_annot(rect, text=REDACT_LABEL, fill=(0, 0, 0))
                if rects:
                    all_findings.append(finding)
                else:
                    # Text found by Presidio but not locatable on page (OCR artefact)
                    finding["note"] = "text_not_locatable_on_page"
                    all_findings.append(finding)

            page.apply_redactions()

        passes_log.append({
            "pass_number": pass_num,
            "detections_count": len(all_findings),
            "detections": all_findings,
        })

        doc.save(str(output_path), garbage=4, deflate=True)
        doc.close()

        # Oscillation check
        if pass_num >= 2:
            prev_count: int = passes_log[pass_num - 2]["detections_count"]
            curr_count: int = passes_log[pass_num - 1]["detections_count"]
            if curr_count == prev_count and curr_count > 0:
                print(
                    f"WARNING: Oscillation detected at pass {pass_num} "
                    f"(count={curr_count} unchanged). Stopping.",
                    file=sys.stderr,
                )
                break

        if len(all_findings) == 0:
            break

        current_path = output_path  # subsequent passes operate on the redacted output

    return passes_log


# ---------------------------------------------------------------------------
# DOCX redaction
# ---------------------------------------------------------------------------

def _collect_docx_paragraphs(doc: Any) -> list[Any]:
    """Collect all paragraphs from body, tables, and headers/footers."""
    import docx as python_docx  # type: ignore[import-untyped] # noqa: F401 (imported for type)

    paragraphs: list[Any] = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.extend(cell.paragraphs)
    for section in doc.sections:
        paragraphs.extend(section.header.paragraphs)
        paragraphs.extend(section.footer.paragraphs)
    return paragraphs


def redact_docx(
    input_path: Path,
    output_path: Path,
    analyzer: Any,
    min_confidence: float,
    whitelist_entities: list[str],
    adversarial_passes: int,
) -> list[dict[str, Any]]:
    import docx as python_docx  # type: ignore[import-untyped]

    passes_log: list[dict[str, Any]] = []
    current_path = input_path

    for pass_num in range(1, adversarial_passes + 1):
        doc = python_docx.Document(str(current_path))
        all_findings: list[dict[str, Any]] = []
        paragraphs = _collect_docx_paragraphs(doc)

        for para in paragraphs:
            full_para_text: str = para.text
            if not full_para_text.strip():
                continue

            findings = analyze_text(analyzer, full_para_text, min_confidence, whitelist_entities)
            if not findings:
                continue

            # Replace found text in each run
            for run in para.runs:
                run_text: str = run.text
                for finding in findings:
                    matched: str = finding["matched_text"]
                    if matched in run_text:
                        run_text = run_text.replace(matched, REDACT_LABEL)
                if run_text != run.text:
                    run.text = run_text

            for finding in findings:
                finding["paragraph_preview"] = full_para_text[:60]
                all_findings.append(finding)

        passes_log.append({
            "pass_number": pass_num,
            "detections_count": len(all_findings),
            "detections": all_findings,
        })

        doc.save(str(output_path))

        # Oscillation check
        if pass_num >= 2:
            prev_count = passes_log[pass_num - 2]["detections_count"]
            curr_count = passes_log[pass_num - 1]["detections_count"]
            if curr_count == prev_count and curr_count > 0:
                print(
                    f"WARNING: Oscillation detected at pass {pass_num} "
                    f"(count={curr_count} unchanged). Stopping.",
                    file=sys.stderr,
                )
                break

        if len(all_findings) == 0:
            break

        current_path = output_path

    return passes_log


# ---------------------------------------------------------------------------
# Text extraction helpers
# ---------------------------------------------------------------------------

def extract_text_from_pdf(path: Path) -> str:
    import fitz  # type: ignore[import-untyped]
    doc = fitz.open(str(path))
    texts: list[str] = [str(doc[i].get_text("text")) for i in range(doc.page_count)]
    doc.close()
    return "\n".join(texts)


def extract_text_from_docx(path: Path) -> str:
    import docx as python_docx  # type: ignore[import-untyped]
    doc = python_docx.Document(str(path))
    parts: list[str] = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.extend(p.text for p in cell.paragraphs)
    return "\n".join(parts)


def extract_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in (".docx", ".doc"):
        return extract_text_from_docx(path)
    return extract_text_from_pdf(path)


# ---------------------------------------------------------------------------
# Summary builder
# ---------------------------------------------------------------------------

def build_summary(passes_log: list[dict[str, Any]]) -> dict[str, Any]:
    if not passes_log:
        return {
            "total_redactions": 0,
            "categories": {},
            "adversarial_passes": 0,
            "final_leakage_count": 0,
        }

    first_pass = passes_log[0]
    cats: dict[str, int] = {}
    for d in first_pass["detections"]:
        cats[d["category"]] = cats.get(d["category"], 0) + 1

    return {
        "total_redactions": first_pass["detections_count"],
        "categories": cats,
        "adversarial_passes": len(passes_log),
        "final_leakage_count": passes_log[-1]["detections_count"],
    }


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main() -> int:
    parser = argparse.ArgumentParser(
        description="Redact PII from Australian trust deeds (Privacy Act 1988 Cth s6)."
    )
    parser.add_argument("--input", required=True, help="Path to input PDF or DOCX")
    parser.add_argument("--output", help="Path for redacted output (default: <input>_redacted.<ext>)")
    parser.add_argument("--log", help="Path for audit log JSON (default: <input>_redact_log.json)")
    parser.add_argument("--quality-report", help="Path for OCR quality report JSON")
    parser.add_argument("--adversarial-passes", type=int, default=3, help="Max re-detection passes")
    parser.add_argument("--dry-run", action="store_true", help="Detect only, do not write output")
    parser.add_argument("--strict", action="store_true", help="Exit 1 if leakage remains after passes")
    parser.add_argument("--check-ocr-only", action="store_true", help="OCR quality check only, no redaction")
    parser.add_argument("--no-ner", action="store_true", help="Disable spaCy NER (regex only)")
    parser.add_argument("--min-confidence", type=float, default=0.6, help="Minimum Presidio score (default 0.6)")
    parser.add_argument(
        "--whitelist-entity",
        action="append",
        default=[],
        dest="whitelist_entities",
        help="Entity text to whitelist from redaction (repeatable)",
    )

    args = parser.parse_args()

    input_path = Path(args.input).expanduser().resolve()
    if not input_path.exists():
        print(f"ERROR: Input file not found: {input_path}", file=sys.stderr)
        return 1

    suffix = input_path.suffix.lower()
    if suffix not in (".pdf", ".docx", ".doc"):
        print(f"ERROR: Unsupported file type '{suffix}'. Supported: .pdf, .docx", file=sys.stderr)
        return 1

    # Dependency check
    missing = check_deps(no_ner=args.no_ner)
    if missing:
        print("ERROR: Missing required packages. Install with:\n" + "\n".join(missing), file=sys.stderr)
        return 1

    # Derive output paths
    output_path = (
        Path(args.output).expanduser().resolve()
        if args.output
        else input_path.parent / f"{input_path.stem}_redacted{suffix}"
    )
    log_path = (
        Path(args.log).expanduser().resolve()
        if args.log
        else input_path.parent / f"{input_path.stem}_redact_log.json"
    )
    quality_report_path = (
        Path(args.quality_report).expanduser().resolve()
        if args.quality_report
        else input_path.parent / f"{input_path.stem}_ocr_quality.json"
    )

    # ------------------------------------------------------------------
    # Phase 0: OCR quality assessment
    # ------------------------------------------------------------------
    print(f"[Phase 0] Assessing OCR quality: {input_path.name}", file=sys.stderr)
    quality = assess_ocr_quality(input_path)
    quality_report_path.write_text(json.dumps(quality, indent=2))
    print(
        f"  OCR quality: {quality['ocr_quality']} "
        f"(score={quality['ocr_confidence_score']:.2f}, "
        f"pages_with_text={quality['pages_with_text']}/{quality['total_pages']})",
        file=sys.stderr,
    )

    if args.check_ocr_only:
        print(json.dumps(quality, indent=2))
        return 0

    # Halt on unreadable
    if quality["ocr_quality"] == "unreadable":
        print(
            "\nREDACTION HALTED - DOCUMENT UNREADABLE\n"
            "======================================\n"
            f"OCR confidence score: {quality['ocr_confidence_score']:.2f}\n"
            "This document's text layer is too poor to support reliable automated PII detection.\n\n"
            "Options:\n"
            "A. Re-scan at 300+ DPI and retry.\n"
            "B. Run OCR preprocessing (marker-pdf or tesseract) and retry on the searchable PDF.\n"
            "C. Manual redaction.\n\n"
            "Do NOT distribute the output — PII leakage risk is HIGH.",
            file=sys.stderr,
        )
        return 1

    # Low-confidence warning
    if quality["ocr_quality"] == "low_confidence":
        print(
            "\nOCR QUALITY FLAG — HUMAN REVIEW REQUIRED\n"
            "=========================================\n"
            f"OCR confidence score: {quality['ocr_confidence_score']:.2f} "
            f"({quality['pages_with_text']}/{quality['total_pages']} pages readable)\n"
            f"Empty pages: {quality['empty_pages']}\n"
            "ACTION: Manually verify empty/poor-quality pages before distributing the redacted output.\n",
            file=sys.stderr,
        )

    # ------------------------------------------------------------------
    # Trust name extraction (flagged, NOT auto-redacted)
    # ------------------------------------------------------------------
    print("[Phase 1] Extracting text for trust name detection...", file=sys.stderr)
    full_text = extract_text(input_path)
    trust_names = extract_trust_names(full_text)
    if trust_names:
        print(
            f"  Trust names detected (NOT auto-redacted — require human decision): {trust_names}",
            file=sys.stderr,
        )

    # ------------------------------------------------------------------
    # Dry-run: detect only, no output written
    # ------------------------------------------------------------------
    if args.dry_run:
        print("[Phase 2] Dry-run detection (no output written)...", file=sys.stderr)
        analyzer = build_analyzer(no_ner=args.no_ner, min_confidence=args.min_confidence)
        findings = analyze_text(analyzer, full_text, args.min_confidence, args.whitelist_entities)
        passes_log: list[dict[str, Any]] = [
            {"pass_number": 1, "detections_count": len(findings), "detections": findings}
        ]
        audit: dict[str, Any] = {
            "input_file": str(input_path),
            "output_file": None,
            "timestamp": datetime.now(timezone.utc).isoformat(),
            "dry_run": True,
            **quality,
            "passes": passes_log,
            "trust_name_review_items": trust_names,
            "vision_only_detections": [],
            "summary": build_summary(passes_log),
        }
        print(json.dumps(audit, indent=2))
        return 0

    # ------------------------------------------------------------------
    # Phase 2: Build analyzer
    # ------------------------------------------------------------------
    print("[Phase 2] Building Presidio analyzer...", file=sys.stderr)
    analyzer = build_analyzer(no_ner=args.no_ner, min_confidence=args.min_confidence)

    # ------------------------------------------------------------------
    # Phase 3: Redact with adversarial loop
    # ------------------------------------------------------------------
    print(f"[Phase 3] Redacting — {args.adversarial_passes} adversarial pass(es)...", file=sys.stderr)
    try:
        if suffix in (".docx", ".doc"):
            passes_log = redact_docx(
                input_path, output_path, analyzer,
                args.min_confidence, args.whitelist_entities, args.adversarial_passes,
            )
        else:
            passes_log = redact_pdf(
                input_path, output_path, analyzer,
                args.min_confidence, args.whitelist_entities, args.adversarial_passes,
            )
    except Exception as exc:
        print(f"ERROR during redaction: {exc}", file=sys.stderr)
        traceback.print_exc()
        return 1

    for p in passes_log:
        print(f"  Pass {p['pass_number']}: {p['detections_count']} detection(s)", file=sys.stderr)

    summary = build_summary(passes_log)

    # ------------------------------------------------------------------
    # Phase 4: Strip PDF metadata
    # ------------------------------------------------------------------
    if suffix == ".pdf" and output_path.exists():
        try:
            import fitz  # type: ignore[import-untyped]
            meta_doc = fitz.open(str(output_path))
            meta_doc.set_metadata({})
            meta_doc.save(str(output_path), incremental=False, garbage=4, deflate=True)
            meta_doc.close()
            print("[Phase 4] PDF metadata stripped.", file=sys.stderr)
        except Exception as exc:
            print(f"WARNING: Could not strip PDF metadata: {exc}", file=sys.stderr)

    # ------------------------------------------------------------------
    # Write audit log
    # ------------------------------------------------------------------
    audit = {
        "input_file": str(input_path),
        "output_file": str(output_path),
        "timestamp": datetime.now(timezone.utc).isoformat(),
        **quality,
        "passes": passes_log,
        "trust_name_review_items": trust_names,
        "vision_only_detections": [],
        "summary": summary,
    }
    log_path.write_text(json.dumps(audit, indent=2))

    # ------------------------------------------------------------------
    # Structured stdout output (for copilot callers)
    # ------------------------------------------------------------------
    copilot_output: dict[str, Any] = {
        "output_file": str(output_path),
        "audit_log": str(log_path),
        "ocr_quality_report": str(quality_report_path),
        "summary": {
            **summary,
            "ocr_quality": quality["ocr_quality"],
            "ocr_confidence_score": quality["ocr_confidence_score"],
        },
    }
    print(json.dumps(copilot_output, indent=2))

    # ------------------------------------------------------------------
    # Strict mode
    # ------------------------------------------------------------------
    if args.strict and summary["final_leakage_count"] > 0:
        print(
            f"\nSTRICT MODE FAILURE: {summary['final_leakage_count']} PII item(s) "
            "still detected after all passes.",
            file=sys.stderr,
        )
        return 1

    print(
        f"\nDone. Redacted file: {output_path}\n"
        f"Redactions: {summary['total_redactions']} | "
        f"Final leakage: {summary['final_leakage_count']} | "
        f"OCR quality: {quality['ocr_quality']}",
        file=sys.stderr,
    )
    return 0


if __name__ == "__main__":
    sys.exit(main())
